"""
agent_harness/document_capabilities.py

The "create a document about a topic, with images and charts" pipeline —
this is a genuinely different capability from pdf_generate (which exports
THIS CHAT SESSION's transcript). This is closer to what happens when a
person asks an AI assistant to write and produce a fresh document:
research the topic, generate content, find/produce visuals, lay it out,
render to a real file, and validate it before calling it done.

Each step here is a separate, composable capability (Section 16's
"do not create one do_everything() mega-tool" principle still applies) —
dynamic_harness.py chains them via ordinary dependencies, exactly like
every other capability.

HONESTY ABOUT WHAT'S VERIFIED VS ASSUMED
------------------------------------------
- topic_research, chart_generate, document_compose: built entirely on
  things already confirmed working elsewhere in this project (mcp_tools
  dispatch, self.llm.call, matplotlib is a common enough dependency but
  NOT confirmed installed — see _render_chart_png).
- image_search: UNVERIFIED whether mcp_tools.dispatch("tavily_search", ...)
  actually supports an `include_images` param or returns an "images" key
  — I don't have mcp_tools.py's tavily_search implementation in front of
  me. Registered as optional_by_default=True specifically so a wrong
  guess here degrades to "document without images" rather than failing
  the whole pipeline — check your terminal for
  "[DOC] image search returned no images" to know if this needs fixing.
- document_render (PDF path): reuses weasyprint, which mcp_pdf_export.py
  already treats as an optional dependency (HTML fallback exists there
  for exactly the "not installed" case) — same assumption here.
- document_render (DOCX path): needs python-docx (`pip install
  python-docx`), not otherwise used elsewhere in this project as far as
  I've seen — flagged as a new dependency if you want the Word-doc path.
"""
import asyncio
import json
import os
import re
import uuid
from typing import Any, Dict, List, Optional

from . import tools
from .task_state import AgentResult

_OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "generated_documents")
_CHART_DIR = os.path.join(_OUTPUT_DIR, "charts")


def _slug(text: str) -> str:
    s = re.sub(r"[^a-zA-Z0-9]+", "_", text or "document").strip("_").lower()
    return (s or "document")[:60]


def _extract_json(text: str) -> Optional[Dict[str, Any]]:
    text = (text or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    m = re.search(r"\{.*\}", text, re.DOTALL)
    if not m:
        return None
    try:
        return json.loads(m.group(0))
    except json.JSONDecodeError:
        return None


# ═══════════════════════════════════════════════════════════════════════
# RESEARCH — generic topic research (not limited to agriculture; for
# agriculture-specific factual questions, prefer agriculture_rag, which
# is grounded in the trusted knowledge base + claim verification).
# ═══════════════════════════════════════════════════════════════════════

async def h_topic_research(ctx, harness=None, parent_agent_id=None, **args) -> AgentResult:
    topic = args.get("topic") or args.get("query", "")
    if not topic.strip():
        return AgentResult(status="failure", issues=["no topic provided"])

    evidence: List[Dict[str, Any]] = []
    try:
        web_out = await tools.run_mcp_tool("tavily_search", {"query": topic})
        evidence = (web_out or {}).get("results") or []
    except Exception as e:
        print(f"[DOC] web research for topic_research failed ({e!r}) — writing from general knowledge only.")

    snippet_block = "\n\n".join(
        f"[{i+1}] {r.get('title','')}\n{(r.get('content') or r.get('snippet') or '')[:500]}"
        for i, r in enumerate(evidence[:8])
    )
    system = (
        "You are a research writer producing accurate, well-organized content for a "
        "professional document. Write clear prose organized into sections with headings "
        "(use markdown ## for headings). Use the evidence provided plus well-established "
        "general knowledge — do NOT fabricate specific statistics, dates, or claims you "
        "aren't confident about. Cite evidence with [N] tags matching the numbered sources "
        "where you actually used them; do not cite a source for a claim it doesn't support."
    )
    user = (f"TOPIC: {topic}\n\nEVIDENCE:\n{snippet_block or '(no web evidence retrieved)'}\n\n"
            f"Write the content now, organized into clear sections.")
    try:
        content, _usage = await asyncio.to_thread(ctx.llm.call, system, user, 2200, 0.3, None)
    except Exception as e:
        return AgentResult(status="failure", issues=[f"research generation failed: {e}"])

    if not content or not content.strip():
        return AgentResult(status="failure", issues=["empty research content"])
    return AgentResult(status="success", output=content.strip(),
                        confidence=0.7 if evidence else 0.4,
                        evidence=evidence, metadata={"topic": topic})


def verify_topic_research(result: AgentResult, constraints: Dict[str, Any]) -> Dict[str, Any]:
    if result.status != "success":
        return {"ok": False, "issues": result.issues or ["research failed"]}
    text = str(result.output or "")
    issues = []
    if len(text) < 200:
        issues.append("research content suspiciously short")
    return {"ok": not issues, "issues": issues}


# ═══════════════════════════════════════════════════════════════════════
# IMAGES — see module docstring: this is the unverified piece.
# ═══════════════════════════════════════════════════════════════════════

async def h_image_search(ctx, **args) -> AgentResult:
    query = args.get("query") or args.get("topic", "")
    try:
        out = await tools.run_mcp_tool("tavily_search", {"query": query, "include_images": True})
    except Exception as e:
        return AgentResult(status="partial", issues=[f"image search unavailable: {e}"], output=[])
    images = (out or {}).get("images") or []
    if not images:
        print("[DOC] image search returned no images — proceeding without images "
              "(check mcp_tools.dispatch's tavily_search 'include_images' support if this always happens).")
        return AgentResult(status="partial", issues=["no images found"], output=[])
    return AgentResult(status="success", output=images[:6], confidence=0.6)


# ═══════════════════════════════════════════════════════════════════════
# CHARTS — LLM proposes ONE chart spec only if the content actually has
# chartable data (never invents numbers); matplotlib renders it.
# ═══════════════════════════════════════════════════════════════════════

def _render_chart_png(spec: Dict[str, Any]) -> str:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    os.makedirs(_CHART_DIR, exist_ok=True)
    labels = spec.get("labels", [])
    values = spec.get("values", [])
    kind = spec.get("chart_type", "bar")

    fig, ax = plt.subplots(figsize=(6, 4))
    if kind == "pie":
        ax.pie(values, labels=labels, autopct="%1.0f%%")
    elif kind == "line":
        ax.plot(labels, values, marker="o")
    else:
        ax.bar(labels, values)
    ax.set_title(spec.get("title", ""))
    fig.tight_layout()

    path = os.path.join(_CHART_DIR, f"chart_{uuid.uuid4().hex[:8]}.png")
    fig.savefig(path, dpi=120)
    plt.close(fig)
    return path


async def h_chart_generate(ctx, **args) -> AgentResult:
    topic = args.get("topic", "")
    research_text = args.get("research_text", "")
    system = (
        "You produce chart specifications for a document generator. Given a topic and "
        "researched content, propose ONE simple, genuinely useful chart ONLY if the "
        "content contains real comparable numeric data (yields, percentages, counts, "
        "trends). If nothing in the content is actually chartable, say so plainly — "
        "never invent numbers to make a chart possible.\n\n"
        "Respond with ONLY this JSON:\n"
        '{"chartable": true|false, "chart_type": "bar"|"line"|"pie", "title": "...", '
        '"labels": ["..."], "values": [<numbers, same length as labels>], '
        '"reason": "<=15 words, only if chartable is false"}'
    )
    user = f"TOPIC: {topic}\n\nCONTENT:\n{research_text[:3000]}\n\nJSON:"
    try:
        raw, _usage = await asyncio.to_thread(ctx.llm.call, system, user, 400, 0.0, None)
    except Exception as e:
        return AgentResult(status="partial", issues=[f"chart spec generation failed: {e}"])

    spec = _extract_json(raw)
    if not spec or not spec.get("chartable"):
        reason = (spec or {}).get("reason", "no chartable data found in the researched content")
        return AgentResult(status="partial", issues=[reason])

    try:
        path = await asyncio.to_thread(_render_chart_png, spec)
    except ImportError:
        return AgentResult(status="partial", issues=["matplotlib not installed — skipping chart"])
    except Exception as e:
        return AgentResult(status="partial", issues=[f"chart rendering failed: {e}"])

    return AgentResult(status="success", output={"path": path, "title": spec.get("title", "")}, confidence=0.6)


# ═══════════════════════════════════════════════════════════════════════
# COMPOSE — assembles researched text + images + chart into one HTML doc.
# Structure/Layout/Content are deliberately collapsed into one step for
# this first pass (spec's ResearchAgent/StructureAgent/LayoutAgent
# distinction is a further refinement, not implemented separately yet).
# ═══════════════════════════════════════════════════════════════════════

_DOC_CSS = """
body { font-family: 'Segoe UI', Helvetica, Arial, sans-serif; color:#222;
       max-width:800px; margin:40px auto; line-height:1.65; padding:0 20px; }
h1 { color:#2a6a0a; border-bottom:3px solid #7ab648; padding-bottom:8px; }
h2 { color:#3a6020; margin-top:32px; }
img { border-radius:8px; box-shadow:0 2px 8px rgba(0,0,0,0.15); }
.figure { margin:20px 0; text-align:center; }
.figure img { max-width:100%; }
.figure .caption { font-size:13px; color:#666; margin-top:6px; }
.sources { font-size:13px; color:#555; margin-top:36px; border-top:1px solid #ddd; padding-top:16px; }
.sources li { margin-bottom:4px; }
"""


def _markdown_headings_to_html(text: str) -> str:
    """Minimal, dependency-free ## heading -> <h2> conversion — the
    research prompt asks for markdown ## headings specifically so this
    conversion has something reliable to match against."""
    lines = []
    for line in text.split("\n"):
        if line.startswith("## "):
            lines.append(f"<h2>{line[3:].strip()}</h2>")
        elif line.strip():
            lines.append(f"<p>{line.strip()}</p>")
    return "\n".join(lines)


async def h_document_compose(ctx, **args) -> AgentResult:
    title = args.get("title") or args.get("topic") or "Document"
    research_text = args.get("research_text", "")
    images = args.get("images") or []
    chart = args.get("chart")
    sources = args.get("sources") or []

    if not research_text.strip():
        return AgentResult(status="failure", issues=["no content to compose — upstream research produced nothing"])

    body_html = _markdown_headings_to_html(research_text)

    figures_html = ""
    for img in images[:4]:
        url = img.get("url") if isinstance(img, dict) else None
        if url:
            figures_html += f'<div class="figure"><img src="{url}" /></div>'

    chart_html = ""
    if chart and chart.get("path"):
        chart_html = (f'<div class="figure"><img src="file://{chart["path"]}" />'
                      f'<div class="caption">{chart.get("title","")}</div></div>')

    sources_html = "".join(
        f"<li>[{i+1}] {s.get('title') or s.get('label') or s.get('source_file','')}</li>"
        for i, s in enumerate(sources)
    )
    sources_block = (f'<div class="sources"><h2>Sources</h2><ul>{sources_html}</ul></div>'
                     if sources_html else "")

    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>{_DOC_CSS}</style></head>
<body>
<h1>{title}</h1>
{figures_html}
{body_html}
{chart_html}
{sources_block}
</body></html>"""

    return AgentResult(status="success", output=html, confidence=1.0, metadata={"title": title})


def verify_document_compose(result: AgentResult, constraints: Dict[str, Any]) -> Dict[str, Any]:
    if result.status != "success":
        return {"ok": False, "issues": result.issues or ["compose failed"]}
    html = str(result.output or "")
    issues = []
    if "<h1>" not in html:
        issues.append("composed document missing a title heading")
    if len(html) < 300:
        issues.append("composed document suspiciously short")
    return {"ok": not issues, "issues": issues}


# ═══════════════════════════════════════════════════════════════════════
# RENDER — HTML -> real PDF (weasyprint) or DOCX (python-docx).
# ═══════════════════════════════════════════════════════════════════════

def _render_pdf(html: str, title: str):
    os.makedirs(_OUTPUT_DIR, exist_ok=True)
    path = os.path.join(_OUTPUT_DIR, f"{_slug(title)}_{uuid.uuid4().hex[:8]}.pdf")
    try:
        from weasyprint import HTML
    except ImportError:
        raise RuntimeError(
            "weasyprint not installed — cannot render PDF (same optional dependency "
            "mcp_pdf_export.py already falls back from for its own HTML-fallback path). "
            "pip install weasyprint, or request doc_format='docx' instead.")
    HTML(string=html).write_pdf(path)
    with open(path, "rb") as f:
        data = f.read()
    return data, path


def _render_docx(html: str, title: str) -> str:
    os.makedirs(_OUTPUT_DIR, exist_ok=True)
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx not installed — pip install python-docx, "
                           "or request doc_format='pdf' instead.")
    # Crude HTML->text: this is a deliberate simplification, not a full
    # HTML-to-DOCX converter — headings/paragraphs survive, inline images
    # and styling do not. Good enough for a first pass; a real DOCX
    # layout pass (matching image/chart placement) is future work.
    text = re.sub(r"<style>.*?</style>", "", html, flags=re.DOTALL)
    text = re.sub(r"<h1>(.*?)</h1>", r"\n# \1\n", text)
    text = re.sub(r"<h2>(.*?)</h2>", r"\n## \1\n", text)
    text = re.sub(r"<[^>]+>", "\n", text)

    doc = Document()
    doc.add_heading(title, level=1)
    for para in [p.strip() for p in text.split("\n") if p.strip() and not p.strip().startswith("#")]:
        doc.add_paragraph(para)

    path = os.path.join(_OUTPUT_DIR, f"{_slug(title)}_{uuid.uuid4().hex[:8]}.docx")
    doc.save(path)
    return path


async def h_document_render(ctx, **args) -> AgentResult:
    html = args.get("html", "")
    doc_format = (args.get("doc_format") or "pdf").lower()
    title = args.get("title") or "document"

    if not html:
        return AgentResult(status="failure", issues=["no composed HTML to render"])

    if doc_format == "docx":
        try:
            path = await asyncio.to_thread(_render_docx, html, title)
        except Exception as e:
            return AgentResult(status="failure", issues=[str(e)])
        return AgentResult(status="success", confidence=1.0,
                            output={"type": "docx", "path": path, "filename": os.path.basename(path)})

    try:
        pdf_bytes, path = await asyncio.to_thread(_render_pdf, html, title)
    except Exception as e:
        return AgentResult(status="failure", issues=[str(e)])
    return AgentResult(status="success", confidence=1.0,
                        output={"type": "pdf", "bytes": pdf_bytes, "path": path, "filename": os.path.basename(path)})


def verify_document_render(result: AgentResult, constraints: Dict[str, Any]) -> Dict[str, Any]:
    if result.status != "success":
        return {"ok": False, "issues": result.issues or ["render failed"]}
    out = result.output or {}
    if not out.get("path") or not os.path.exists(out["path"]):
        return {"ok": False, "issues": ["rendered file does not exist on disk"]}
    return {"ok": True, "issues": []}


# ═══════════════════════════════════════════════════════════════════════
# VALIDATE — integrity check on the actual rendered file.
# ═══════════════════════════════════════════════════════════════════════

async def h_document_validate(ctx, **args) -> AgentResult:
    render = args.get("render_result") or {}
    if render.get("type") == "pdf":
        pdf_bytes = render.get("bytes", b"")
        ok = len(pdf_bytes) > 500
        page_count = None
        try:
            try:
                import pymupdf as fitz
            except ImportError:
                import fitz
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            page_count = doc.page_count
            ok = ok and page_count > 0
        except ImportError:
            pass
        except Exception as e:
            ok = False
            return AgentResult(status="failure", issues=[f"PDF failed to open: {e}"])
        result = {"ok": ok, "type": "pdf", "size_bytes": len(pdf_bytes), "page_count": page_count}
        return AgentResult(status="success" if ok else "failure", output=result,
                            confidence=1.0 if ok else 0.0,
                            issues=[] if ok else ["PDF failed integrity check"])

    if render.get("type") == "docx":
        path = render.get("path", "")
        ok = bool(path) and os.path.exists(path) and os.path.getsize(path) > 500
        return AgentResult(status="success" if ok else "failure",
                            output={"ok": ok, "type": "docx", "size_bytes": os.path.getsize(path) if ok else 0},
                            confidence=1.0 if ok else 0.0,
                            issues=[] if ok else ["docx file missing or too small"])

    return AgentResult(status="failure", issues=["unexpected render result — nothing to validate"])
